from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Recent Platform Enhancements & Feature Updates', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    features = [
        {
            "title": "1. Multi-Audience WhatsApp & AI Chatbot System",
            "update": "We have launched a specialized AI-powered assistant for Merchants, running in parallel with our Customer chatbot.",
            "how": "The new Merchant bot is 'business-aware.' It can answer specific questions about recent orders, pending payouts, and real-time inventory levels directly on WhatsApp. We also implemented a new secure phone-linking system for merchants to receive automated business notifications.",
            "benefit": "Merchants get instant, 24/7 access to their business health without needing to log into the dashboard, while customers benefit from a more responsive communication ecosystem."
        },
        {
            "title": "2. Dynamic 'Mera Paisa' (Growth Fund) Ecosystem",
            "update": "Full implementation of the 'Mera Paisa' investment platform, providing a robust growth engine for the merchant community.",
            "how": "We’ve added dedicated dashboards for merchants to track growth fund balances and returns, along with powerful administrative tools to manage investment orders and monitor fund health.",
            "benefit": "Empowers merchants to reinvest and grow their capital within the platform's ecosystem, offering a professional financial growth tool built directly into their daily workflow."
        },
        {
            "title": "3. 'Ironclad' Rewards & Wallet Integration",
            "update": "The 'InTrust Rewards' engine has been hardened and expanded to support all payment methods.",
            "how": "Users now earn reward points even when paying via their digital wallet for gift cards or shopping. We also implemented advanced 'idempotency' guards that ensure rewards are credited accurately exactly once, even if there are network glitches or duplicate payment signals.",
            "benefit": "This ensures a perfectly reliable reward experience, building deep trust with users and encouraging higher wallet usage and platform spend."
        },
        {
            "title": "4. Optimized Daily User Engagement",
            "update": "The daily reward loop has been simplified to maximize user satisfaction.",
            "how": "We have standardized the daily bonus to a flat 50 points and enabled Instant Claims, removing the 'scratch card' requirement for this specific action to speed up the user experience.",
            "benefit": "Users can collect their bonus in a single click, making daily platform engagement faster and more rewarding."
        },
        {
            "title": "5. Premium UI & Visual Modernization",
            "update": "Significant design overhaul of the Shop and Wholesale interfaces.",
            "how": "We integrated dynamic brand marquees, interactive rating systems, and high-performance layouts that provide a 'premium mall' feel to the e-commerce experience.",
            "benefit": "Elevates the platform's brand perception to international standards, providing a world-class experience for both buyers and sellers."
        }
    ]
    
    for feature in features:
        doc.add_heading(feature['title'], level=1)
        
        p = doc.add_paragraph()
        p.add_run('The Update: ').bold = True
        p.add_run(feature['update'])
        
        p = doc.add_paragraph()
        p.add_run('How it Works: ').bold = True
        p.add_run(feature['how'])
        
        p = doc.add_paragraph()
        p.add_run('The Benefit: ').bold = True
        p.add_run(feature['benefit'])
        
        doc.add_paragraph() # Spacer

    doc.save('Project_Updates_May_2026.docx')
    print("Document created successfully: Project_Updates_May_2026.docx")

if __name__ == "__main__":
    create_document()
